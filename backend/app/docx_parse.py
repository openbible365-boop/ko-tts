"""解析上传的 .docx 范文 -> 待录行列表。

规则: 取正文段落(document.paragraphs, 忽略表格/图片/页眉页脚), 段落按回车分行;
段内若含多句(整篇连写成一段的常见情况), 再按句末标点「。！？!?」切句。
这样无论 docx 是"一句一段"还是"整段连写", 都能拆成多行。空行/空白跳过。

纯函数 parse_docx_lines(bytes) -> list[str], 便于单测。
"""

import io
import re

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


class DocxParseError(ValueError):
    """上传的文件不是合法的 .docx (损坏或非 Word 格式)。"""


# 一句 = 直到句末标点(中英)为止, 句末的右引号/右括号并入本句; 末尾无标点的残段也算一句
_SENTENCE_RE = re.compile(r'[^。！？!?]*[。！？!?]+[”"’\'』」）)】\]]*|[^。！？!?]+')


def _split_sentences(text: str) -> list[str]:
    out: list[str] = []
    for m in _SENTENCE_RE.finditer(text):
        s = m.group().strip().strip("　").strip()  # 去普通空白与全角空格
        if s:
            out.append(s)
    return out


def parse_docx_lines(data: bytes) -> list[str]:
    try:
        doc = Document(io.BytesIO(data))
    except PackageNotFoundError as e:
        raise DocxParseError("无法解析文件: 不是合法的 .docx 文档") from e
    except Exception as e:  # 损坏的 zip / 底层格式异常, 一律当作不可解析
        raise DocxParseError("无法解析文件: 文档损坏或格式不受支持") from e

    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.extend(_split_sentences(text))
    return lines
